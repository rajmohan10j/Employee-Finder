import os
import re
import json
import asyncio
import urllib.parse
from pathlib import Path
from docx import Document
from playwright.async_api import async_playwright
from google import genai

OUTPUT_DIR = Path("./candidate_data")
RESUMES_DIR = OUTPUT_DIR / "resumes"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
RESUMES_DIR.mkdir(parents=True, exist_ok=True)

# Public sample resumes used as fallbacks if Google Search blocks automated requests
FALLBACK_SAMPLES = [
    "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf",
    "https://www.orimi.com/pdf-test.pdf"
]

async def extract_candidate_info(file_path: Path, client: genai.Client) -> dict:
    print(f"   🤖 Sending resume to Gemini API for parsing...")
    
    prompt = """
    Analyze this document/resume and extract candidate details in pure JSON format:
    {
      "full_name": "Full name or Candidate_X if not found",
      "dob_or_age": "Age or DOB or Unknown",
      "total_experience": "Years of experience or Unknown",
      "phone_number": "Phone number or Unknown",
      "email": "Email address or Unknown"
    }
    Return ONLY valid JSON with no markdown formatting.
    """
    
    try:
        uploaded_file = client.files.upload(file=str(file_path))
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[uploaded_file, prompt]
        )
        raw_text = response.text.strip().replace("```json", "").replace("```", "").strip()
        data = json.loads(raw_text)
        return data
    except Exception as e:
        print(f"   ⚠️ Gemini extraction notice: {e}")
        return {
            "full_name": file_path.stem,
            "dob_or_age": "Unknown",
            "total_experience": "Unknown",
            "phone_number": "Unknown",
            "email": "Unknown"
        }

def save_docx_report(candidates: list, output_path: Path):
    doc = Document()
    doc.add_heading("Candidate Sourcing Report", level=0)
    
    for c in candidates:
        doc.add_heading(c.get("full_name", "Unknown Candidate"), level=1)
        p = doc.add_paragraph()
        p.add_run("Age / DOB: ").bold = True
        p.add_run(f"{c.get('dob_or_age')}\n")
        p.add_run("Total Experience: ").bold = True
        p.add_run(f"{c.get('total_experience')}\n")
        p.add_run("Phone Number: ").bold = True
        p.add_run(f"{c.get('phone_number')}\n")
        p.add_run("Email: ").bold = True
        p.add_run(f"{c.get('email')}\n")
        p.add_run("Saved Resume File: ").bold = True
        p.add_run(f"{c.get('saved_filename')}\n")
        doc.add_paragraph("-" * 40)
        
    doc.save(str(output_path))
    print(f"📄 Word report successfully saved to: {output_path}")

async def run_automation():
    print("=" * 65)
    print("      AUTOMATED JOB SEEKER RESUME FETCHER & PARSER v2")
    print("=" * 65)
    
    api_key = input("1. Enter your Gemini API Key: ").strip()
    if not api_key:
        print("❌ Error: Gemini API Key is required.")
        return

    client = genai.Client(api_key=api_key)
    
    query = input("2. Enter job query (e.g. 'Software Engineer resume filetype:pdf'): ").strip()
    if not query:
        query = "Python Developer resume filetype:pdf"
        print(f"   Using default query: '{query}'")
        
    print("\n🚀 Launching automated browser session...\n")
    candidates_list = []
    
    async with async_playwright() as p:
        # Launch Chrome with realistic user flags
        browser = await p.chromium.launch(
            headless=False,
            args=["--disable-blink-features=AutomationControlled"]
        )
        context = await browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            accept_downloads=True
        )
        page = await context.new_page()
        
        search_url = f"https://www.google.com/search?q={urllib.parse.quote(query)}"
        print(f"🔍 Navigating to Google Search...")
        await page.goto(search_url, wait_until="domcontentloaded")
        await page.wait_for_timeout(3000)
        
        # Extract external web links from Google results
        hrefs = await page.eval_on_selector_all("a", "elements => elements.map(e => e.href)")
        valid_links = []
        for h in hrefs:
            if h and h.startswith("http") and "google.com" not in h and "youtube.com" not in h:
                if h not in valid_links:
                    valid_links.append(h)
                    
        valid_links = valid_links[:5]
        
        if not valid_links:
            print("⚠️ Google Search blocked automated link extraction or returned 0 links.")
            print("🔄 Falling back to test mode using direct sample document URLs to demonstrate full execution...\n")
            valid_links = FALLBACK_SAMPLES
        else:
            print(f"🔗 Found {len(valid_links)} candidate links to process.\n")
            
        for idx, link in enumerate(valid_links, 1):
            print(f"[{idx}/{len(valid_links)}] Processing link: {link}")
            temp_path = RESUMES_DIR / f"temp_{idx}.pdf"
            
            try:
                # Direct download or fetch content bytes
                response = await page.request.get(link, timeout=10000)
                if response.status == 200:
                    body = await response.body()
                    with open(temp_path, "wb") as f:
                        f.write(body)
                    print(f"   📥 Downloaded file ({len(body)} bytes)")
                else:
                    print(f"   ⚠️ Could not download file directly (Status {response.status})")
                    continue
                    
                # Parse with Gemini API
                info = await extract_candidate_info(temp_path, client)
                
                # Format name and rename file
                raw_name = info.get("full_name", f"Candidate_{idx}")
                clean_name = re.sub(r'[^a-zA-Z0-9_\- ]', '', raw_name).strip().replace(" ", "_")
                if not clean_name:
                    clean_name = f"Candidate_{idx}"
                    
                final_filename = f"{clean_name}.pdf"
                final_file_path = RESUMES_DIR / final_filename
                
                if temp_path.exists():
                    if final_file_path.exists():
                        final_file_path.unlink()
                    temp_path.rename(final_file_path)
                    
                info["saved_filename"] = final_filename
                candidates_list.append(info)
                print(f"   🏷️ Saved resume as: {final_filename}\n")
                
            except Exception as e:
                print(f"   ⏩ Link failed or skipped: {e}\n")
                if temp_path.exists():
                    temp_path.unlink()
                continue
                
        await browser.close()
        
    # Generate final output reports
    if candidates_list:
        json_report = OUTPUT_DIR / "candidates_report.json"
        docx_report = OUTPUT_DIR / "candidates_report.docx"
        
        with open(json_report, "w", encoding="utf-8") as f:
            json.dump(candidates_list, f, indent=2)
            
        print(f"✅ JSON report saved to: {json_report}")
        save_docx_report(candidates_list, docx_report)
        print("\n🎉 AUTOMATION COMPLETE! Open 'candidate_data' folder to view outputs.")
    else:
        print("❌ No candidate files were downloaded.")

if __name__ == "__main__":
    asyncio.run(run_automation())
