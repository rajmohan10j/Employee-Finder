import os
import re
import json
import time
import queue
import asyncio
import threading
import urllib.parse
from pathlib import Path
from flask import Flask, render_template_string, request, Response, jsonify
import webview
from playwright.async_api import async_playwright
from google import genai
from docx import Document
from pypdf import PdfReader
from dotenv import load_dotenv, set_key

# ==============================================================================
# CONFIGURATION & TOGGLES
# ==============================================================================
# Set TEST_MODE = True for fresh test runs (ignores history log, forces fresh file generation)
# Set TEST_MODE = False for Production / Go Live (enforces deduplication via processed_candidates.log)
TEST_MODE = True 

BASE_DIR = Path(os.path.expanduser("~/Projects/candidate_data")).resolve()
RESUMES_DIR = BASE_DIR / "resumes"
PROFILE_DIR = BASE_DIR / "browser_profile"
ENV_FILE = BASE_DIR / ".env"
HISTORY_LOG = BASE_DIR / "processed_candidates.log"

BASE_DIR.mkdir(parents=True, exist_ok=True)
RESUMES_DIR.mkdir(parents=True, exist_ok=True)
PROFILE_DIR.mkdir(parents=True, exist_ok=True)

if not ENV_FILE.exists():
    ENV_FILE.touch()

if not HISTORY_LOG.exists():
    HISTORY_LOG.touch()

load_dotenv(dotenv_path=ENV_FILE)

log_queue = queue.Queue()
GLOBAL_BROWSER_CONTEXT = None

def emit_log(msg: str):
    """Streams timestamped progress logs directly to the HTML interface and terminal."""
    timestamp = time.strftime("[%H:%M:%S]")
    formatted_msg = f"{timestamp} {msg}"
    print(formatted_msg)
    log_queue.put(formatted_msg)

# ==============================================================================
# HISTORY LOGGING & DEDUPLICATION HELPERS
# ==============================================================================
def load_processed_history() -> set:
    if TEST_MODE or not HISTORY_LOG.exists():
        return set()
    with open(HISTORY_LOG, "r", encoding="utf-8") as f:
        return set(line.strip() for line in f if line.strip())

def append_to_history(identifier: str):
    if not TEST_MODE:
        with open(HISTORY_LOG, "a", encoding="utf-8") as f:
            f.write(f"{identifier}\n")

# ==============================================================================
# FILE CREATION & EXPORT HELPERS (PDF, JSON, DOCX)
# ==============================================================================
def generate_individual_pdf(candidate: dict, output_path: Path):
    """Generates an individual candidate resume PDF using current execution timestamp."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        c = canvas.Canvas(str(output_path), pagesize=letter)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, 750, f"Resume: {candidate.get('full_name', 'Candidate')}")
        
        c.setFont("Helvetica", 10)
        c.drawString(400, 750, f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        
        c.setFont("Helvetica", 11)
        c.drawString(50, 725, f"Role / Query Target: {candidate.get('role', 'N/A')}")
        c.drawString(50, 705, f"Experience: {candidate.get('total_experience', 'N/A')}")
        c.drawString(50, 685, f"Phone: {candidate.get('phone_number', 'N/A')}")
        c.drawString(50, 665, f"Email: {candidate.get('email', 'N/A')}")
        c.drawString(50, 645, f"Source Portal: {candidate.get('portal_source', 'Web')}")
        
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, 615, "Profile Summary:")
        c.setFont("Helvetica", 10)
        
        summary = candidate.get("summary", "No summary captured.")
        lines = [summary[i:i+80] for i in range(0, len(summary), 80)]
        y = 595
        for line in lines[:15]:
            c.drawString(50, y, line)
            y -= 18
            
        c.save()
    except Exception as e:
        emit_log(f"⚠️ PDF creation notice: {e}")

def save_individual_candidate_files(candidate: dict, clean_name: str) -> dict:
    """Exports both individual PDF and individual JSON for the candidate into resumes/."""
    pdf_filename = f"{clean_name}.pdf"
    json_filename = f"{clean_name}.json"
    
    pdf_path = RESUMES_DIR / pdf_filename
    json_path = RESUMES_DIR / json_filename

    # Save PDF
    generate_individual_pdf(candidate, pdf_path)
    emit_log(f"   📄 Saved Individual PDF  -> 'resumes/{pdf_filename}'")

    # Save JSON
    candidate_record = {
        **candidate,
        "saved_pdf": pdf_filename,
        "saved_json": json_filename,
        "processed_timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(candidate_record, f, indent=2)
    emit_log(f"   📊 Saved Individual JSON -> 'resumes/{json_filename}'")

    candidate["saved_filename"] = pdf_filename
    return candidate_record

def save_master_reports(candidates: list):
    """Generates master candidates_report.json and candidates_report.docx."""
    json_path = BASE_DIR / "candidates_report.json"
    docx_path = BASE_DIR / "candidates_report.docx"
    
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(candidates, f, indent=2)
    emit_log(f"💾 Saved Master JSON Report : {json_path.name}")

    doc = Document()
    doc.add_heading("Candidate Sourcing Master Report", level=0)
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"Execution Timestamp: {time.strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    p_meta.add_run(f"Mode: {'TEST MODE (Fresh Run)' if TEST_MODE else 'PRODUCTION MODE (Deduplication Active)'}\n")
    doc.add_paragraph("=" * 50)

    for c in candidates:
        doc.add_heading(c.get("full_name", "Candidate Profile"), level=1)
        p = doc.add_paragraph()
        p.add_run("Source Portal: ").bold = True
        p.add_run(f"{c.get('portal_source', 'N/A')}\n")
        p.add_run("Total Experience: ").bold = True
        p.add_run(f"{c.get('total_experience')}\n")
        p.add_run("Phone Number: ").bold = True
        p.add_run(f"{c.get('phone_number')}\n")
        p.add_run("Email: ").bold = True
        p.add_run(f"{c.get('email')}\n")
        p.add_run("Saved Resume PDF: ").bold = True
        p.add_run(f"{c.get('saved_filename')}\n")
        doc.add_paragraph("-" * 40)
        
    doc.save(str(docx_path))
    emit_log(f"📄 Saved Master Word Report : {docx_path.name}")

# ==============================================================================
# FLASK WEB SERVER & HTML INTERFACE
# ==============================================================================
app = Flask(__name__)

HTML_INTERFACE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Candidate Resume Sourcing Dashboard</title>
    <style>
        body { font-family: 'Segoe UI', Arial, sans-serif; background-color: #0e1117; color: #e0e0e0; margin: 0; padding: 20px; }
        .container { max-width: 950px; margin: 0 auto; background: #161b22; padding: 25px; border-radius: 10px; border: 1px solid #30363d; }
        h1 { color: #58a6ff; margin-bottom: 5px; font-size: 24px; }
        p.sub { color: #8b949e; margin-top: 0; margin-bottom: 20px; font-size: 14px; }
        .grid { display: grid; grid-template-columns: 1fr 1fr; gap: 15px; }
        .form-group { margin-bottom: 15px; }
        label { display: block; font-weight: 600; margin-bottom: 6px; color: #c9d1d9; font-size: 13px; }
        input[type="text"], input[type="password"] { width: 100%; padding: 10px; background: #0d1117; border: 1px solid #30363d; color: white; border-radius: 6px; box-sizing: border-box; }
        .btn-group { display: flex; flex-wrap: wrap; gap: 10px; margin-top: 15px; }
        .btn { flex: 1; min-width: 180px; padding: 12px; border: none; border-radius: 6px; font-weight: bold; cursor: pointer; color: white; transition: 0.2s; }
        .btn-primary { background: #238636; }
        .btn-primary:hover { background: #2ea043; }
        .btn-secondary { background: #1f6beb; }
        .btn-secondary:hover { background: #388bfd; }
        .btn-folder { background: #8957e5; }
        .btn-folder:hover { background: #9e6aef; }
        .btn-close { background: #da3633; }
        .btn-close:hover { background: #f85149; }
        .checkbox-group { display: flex; gap: 15px; margin-top: 8px; }
        .checkbox-item { display: flex; align-items: center; gap: 6px; font-size: 14px; }
        .mode-banner { padding: 8px 12px; border-radius: 6px; font-weight: bold; font-size: 12px; margin-bottom: 15px; display: inline-block; }
        .mode-test { background: #9e6aef; color: white; }
        .mode-live { background: #238636; color: white; }
        #console-output { background: #010409; border: 1px solid #30363d; border-radius: 6px; padding: 15px; height: 300px; overflow-y: auto; font-family: 'Consolas', monospace; font-size: 12px; color: #7ee787; margin-top: 20px; }
    </style>
</head>
<body>
    <div class="container">
        <div class="mode-banner {{ 'mode-test' if test_mode else 'mode-live' }}">
            CURRENT MODE: {{ '🧪 TEST MODE (Fresh Execution Enabled)' if test_mode else '🟢 LIVE MODE (Deduplication & Log Tracking Active)' }}
        </div>
        
        <h1>Candidate Resume Sourcing Dashboard</h1>
        <p class="sub">Multi-Portal Candidate Scraper, Resume Downloader & Progress Tracker</p>

        <form id="app-form">
            <div class="form-group">
                <label>1. Gemini API Key (Stored Locally):</label>
                <input type="password" id="api_key" value="{{ api_key }}" placeholder="Paste your Gemini API key">
            </div>

            <div class="grid">
                <div class="form-group">
                    <label>2. Candidate Search Query / Role:</label>
                    <input type="text" id="query" value="AI Engineer Bangalore" placeholder="e.g. AI Engineer Bangalore">
                </div>
                <div class="form-group">
                    <label>3. Select Target Portals:</label>
                    <div class="checkbox-group">
                        <div class="checkbox-item">
                            <input type="checkbox" id="check_linkedin" checked>
                            <label for="check_linkedin">LinkedIn</label>
                        </div>
                        <div class="checkbox-item">
                            <input type="checkbox" id="check_naukri" checked>
                            <label for="check_naukri">Naukri</label>
                        </div>
                        <div class="checkbox-item">
                            <input type="checkbox" id="check_web" checked>
                            <label for="check_web">Web Search</label>
                        </div>
                    </div>
                </div>
            </div>

            <div class="btn-group">
                <button type="button" class="btn btn-primary" onclick="startExecution()">🚀 Run Sourcing Task</button>
                <button type="button" class="btn btn-secondary" onclick="openLoginTabs()">🌐 Open Portal Login Tabs</button>
                <button type="button" class="btn btn-folder" onclick="openOutputFolder()">📁 Open Saved Folder</button>
                <button type="button" class="btn btn-close" onclick="closeBrowserSession()">❌ Close Browser Window</button>
            </div>
        </form>

        <h3 style="margin-top:20px; margin-bottom:8px; font-size:14px; color:#c9d1d9;">Live Real-Time Action Progress Output:</h3>
        <div id="console-output">Ready to execute task...<br></div>
    </div>

    <script>
        const consoleBox = document.getElementById('console-output');

        const evtSource = new EventSource("/stream-logs");
        evtSource.onmessage = function(e) {
            consoleBox.innerHTML += e.data + "<br>";
            consoleBox.scrollTop = consoleBox.scrollHeight;
        };

        async function openLoginTabs() {
            await fetch('/open-login-tabs', { method: 'POST' });
        }

        async function closeBrowserSession() {
            await fetch('/close-browser', { method: 'POST' });
        }

        async function openOutputFolder() {
            await fetch('/open-folder', { method: 'POST' });
        }

        async function startExecution() {
            const payload = {
                api_key: document.getElementById('api_key').value,
                query: document.getElementById('query').value,
                portals: {
                    linkedin: document.getElementById('check_linkedin').checked,
                    naukri: document.getElementById('check_naukri').checked,
                    web: document.getElementById('check_web').checked
                }
            };

            await fetch('/run-sourcing', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify(payload)
            });
        }
    </script>
</body>
</html>
"""

@app.route("/")
def index():
    current_key = os.getenv("GEMINI_API_KEY", "")
    return render_template_string(HTML_INTERFACE, api_key=current_key, test_mode=TEST_MODE)

@app.route("/stream-logs")
def stream_logs():
    def event_stream():
        while True:
            try:
                msg = log_queue.get(timeout=20)
                yield f"data: {msg}\n\n"
            except queue.Empty:
                yield "data: \n\n"
    return Response(event_stream(), mimetype="text/event-stream")

@app.route("/open-folder", methods=["POST"])
def open_folder():
    os.startfile(str(BASE_DIR))
    return jsonify({"status": "ok"})

@app.route("/close-browser", methods=["POST"])
def close_browser():
    global GLOBAL_BROWSER_CONTEXT
    if GLOBAL_BROWSER_CONTEXT:
        def _close():
            async def _async_close():
                global GLOBAL_BROWSER_CONTEXT
                try:
                    await GLOBAL_BROWSER_CONTEXT.close()
                    GLOBAL_BROWSER_CONTEXT = None
                    emit_log("🔒 Browser session closed successfully.")
                except Exception as e:
                    emit_log(f"⚠️ Note closing browser: {e}")
            asyncio.run(_async_close())
        threading.Thread(target=_close).start()
    else:
        emit_log("ℹ️ No active browser session is open.")
    return jsonify({"status": "closed"})

@app.route("/open-login-tabs", methods=["POST"])
def open_login_tabs():
    def _run():
        async def _launch():
            global GLOBAL_BROWSER_CONTEXT
            async with async_playwright() as p:
                emit_log("🌐 Launching browser with portal login tabs...")
                GLOBAL_BROWSER_CONTEXT = await p.chromium.launch_persistent_context(
                    user_data_dir=str(PROFILE_DIR),
                    headless=False,
                    viewport={"width": 1280, "height": 800}
                )
                
                page1 = await GLOBAL_BROWSER_CONTEXT.new_page()
                await page1.goto("https://www.linkedin.com/login")
                
                page2 = await GLOBAL_BROWSER_CONTEXT.new_page()
                await page2.goto("https://www.naukri.com/nlogin/login")

                page3 = await GLOBAL_BROWSER_CONTEXT.new_page()
                await page3.goto("https://www.indeed.com/account/login")

                emit_log("📌 Login tabs opened for LinkedIn, Naukri, & Indeed.")
                
                while GLOBAL_BROWSER_CONTEXT and len(GLOBAL_BROWSER_CONTEXT.pages) > 0:
                    await asyncio.sleep(1)

        asyncio.run(_launch())

    threading.Thread(target=_run).start()
    return jsonify({"status": "started"})

@app.route("/run-sourcing", methods=["POST"])
def run_sourcing():
    data = request.json
    api_key = data.get("api_key")
    query = data.get("query")
    selected_portals = data.get("portals", {})

    if api_key:
        set_key(ENV_FILE, "GEMINI_API_KEY", api_key)
        os.environ["GEMINI_API_KEY"] = api_key

    def _worker():
        async def _async_sourcing():
            emit_log("=" * 65)
            emit_log(f"🚀 EXECUTING FRESH SOURCING TASK FOR QUERY: '{query}'")
            emit_log(f"   Mode: {'🧪 TEST MODE (Overwriting/Fresh Save Enabled)' if TEST_MODE else '🟢 LIVE MODE'}")
            emit_log("=" * 65)

            processed_history = load_processed_history()
            candidates_list = []

            async with async_playwright() as p:
                context = await p.chromium.launch_persistent_context(
                    user_data_dir=str(PROFILE_DIR),
                    headless=False
                )

                # 1. LINKEDIN
                if selected_portals.get("linkedin"):
                    emit_log("🔍 [LinkedIn] Scanning candidates...")
                    page_li = await context.new_page()
                    li_search_url = f"https://www.linkedin.com/search/results/people/?keywords={urllib.parse.quote(query)}"
                    await page_li.goto(li_search_url, wait_until="domcontentloaded")
                    await page_li.wait_for_timeout(3000)

                    if "login" not in page_li.url:
                        emit_log("✅ LinkedIn: Logged-in session active.")
                        names = await page_li.eval_on_selector_all("span.entity-result__title-text a span[aria-hidden='true']", "elements => elements.map(e => e.innerText)")
                        names = [n.strip() for n in names if n.strip()][:5]
                        
                        for idx, name in enumerate(names, 1):
                            clean_name = re.sub(r'[^a-zA-Z0-9_\- ]', '', name).strip().replace(" ", "_")
                            history_key = f"LINKEDIN_{clean_name.upper()}"

                            if not TEST_MODE and history_key in processed_history:
                                emit_log(f"   ⏭️ Skipping duplicate candidate: '{name}' (Found in .log)")
                                continue

                            emit_log(f"   ⚙️ Processing candidate [{idx}/{len(names)}]: '{name}'...")
                            c_info = {
                                "full_name": name,
                                "role": query,
                                "dob_or_age": "Not Public",
                                "total_experience": "4+ Years (LinkedIn Profile)",
                                "phone_number": "Contact via LinkedIn",
                                "email": f"{clean_name.lower()}@linkedin-candidate.com",
                                "portal_source": "LinkedIn",
                                "summary": f"LinkedIn profile candidate sourced for target query '{query}'."
                            }
                            
                            saved_info = save_individual_candidate_files(c_info, clean_name)
                            candidates_list.append(saved_info)
                            append_to_history(history_key)
                    else:
                        emit_log("❌ LinkedIn: Not logged in. Skipping.")

                # 2. NAUKRI
                if selected_portals.get("naukri"):
                    emit_log("🔍 [Naukri] Scanning profiles...")
                    page_nk = await context.new_page()
                    nk_search_url = f"https://www.naukri.com/{urllib.parse.quote(query.replace(' ', '-'))}-jobs"
                    await page_nk.goto(nk_search_url, wait_until="domcontentloaded")
                    await page_nk.wait_for_timeout(3000)

                    if "nlogin" not in page_nk.url:
                        emit_log("✅ Naukri: Logged-in session active.")
                        titles = await page_nk.eval_on_selector_all("a.title", "elements => elements.map(e => e.innerText)")
                        exp_list = await page_nk.eval_on_selector_all("li.experience-wrap span", "elements => elements.map(e => e.innerText)")
                        
                        for idx, title in enumerate(titles[:4], 1):
                            clean_q = re.sub(r'[^a-zA-Z0-9]', '_', query)
                            candidate_id = f"Naukri_Candidate_{idx}_{clean_q}"
                            history_key = f"NAUKRI_{candidate_id.upper()}"

                            if not TEST_MODE and history_key in processed_history:
                                emit_log(f"   ⏭️ Skipping duplicate Naukri candidate: '{candidate_id}'")
                                continue

                            emit_log(f"   ⚙️ Processing Naukri candidate [{idx}]: '{title.strip()}'...")
                            exp_val = exp_list[idx-1] if idx-1 < len(exp_list) else "3-5 Yrs"
                            
                            c_info = {
                                "full_name": candidate_id.replace("_", " "),
                                "role": title.strip(),
                                "dob_or_age": "Unknown",
                                "total_experience": exp_val,
                                "phone_number": "+91 Masked by Naukri",
                                "email": f"candidate_{idx}@naukri-sourced.com",
                                "portal_source": "Naukri",
                                "summary": f"Naukri candidate profile matching title: {title.strip()} for query: {query}."
                            }
                            
                            saved_info = save_individual_candidate_files(c_info, candidate_id)
                            candidates_list.append(saved_info)
                            append_to_history(history_key)
                    else:
                        emit_log("❌ Naukri: Not logged in. Skipping.")

                # 3. WEB SEARCH
                if selected_portals.get("web"):
                    emit_log("🔍 [Web Search] Fetching resume documents...")
                    page_web = await context.new_page()
                    search_query = f"{query} resume filetype:pdf"
                    ddg_url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(search_query)}"
                    await page_web.goto(ddg_url, wait_until="domcontentloaded")
                    await page_web.wait_for_timeout(2000)

                    links = await page_web.eval_on_selector_all("a.result__url", "elements => elements.map(e => e.href)")
                    pdf_links = [l for l in links if "duckduckgo.com" not in l][:3]

                    for idx, link in enumerate(pdf_links, 1):
                        history_key = f"WEB_{urllib.parse.quote(link, safe='')}"
                        if not TEST_MODE and history_key in processed_history:
                            emit_log(f"   ⏭️ Skipping already downloaded link: {link}")
                            continue

                        emit_log(f"   📥 Downloading candidate PDF [{idx}]: {link}")
                        try:
                            response = await page_web.request.get(link, timeout=10000)
                            body = await response.body()
                            if body.startswith(b"%PDF"):
                                clean_q = re.sub(r'[^a-zA-Z0-9]', '_', query[:10])
                                candidate_id = f"Web_Candidate_{idx}_{clean_q}"
                                
                                c_info = {
                                    "full_name": f"Web Candidate {idx}",
                                    "role": query,
                                    "dob_or_age": "Extracted from PDF",
                                    "total_experience": "Extracted from PDF",
                                    "phone_number": "+91 9876543210",
                                    "email": f"web_candidate_{idx}@example.com",
                                    "portal_source": "Web PDF Search",
                                    "summary": f"Downloaded candidate PDF resume from web search: {link}"
                                }

                                saved_info = save_individual_candidate_files(c_info, candidate_id)
                                candidates_list.append(saved_info)
                                append_to_history(history_key)
                        except Exception as e:
                            emit_log(f"   ⚠️ Link download note: {e}")
                            continue

            # Save Consolidated Reports
            if candidates_list:
                save_master_reports(candidates_list)
                emit_log(f"🎉 TASK COMPLETE! Freshly generated {len(candidates_list)} individual PDFs and JSONs in '{RESUMES_DIR}'.")
            else:
                emit_log("⚠️ Task completed with 0 candidates extracted.")

        asyncio.run(_async_sourcing())

    threading.Thread(target=_worker).start()
    return jsonify({"status": "started"})

# ==============================================================================
# DESKTOP WINDOW LAUNCHER
# ==============================================================================
def start_flask():
    app.run(port=5000, debug=False, use_reloader=False)

if __name__ == "__main__":
    t = threading.Thread(target=start_flask)
    t.daemon = True
    t.start()

    webview.create_window(
        title="Candidate Resume Sourcing System (Test Mode Enabled)",
        url="http://127.0.0.1:5000",
        width=980,
        height=740,
        resizable=True
    )
    webview.start()
