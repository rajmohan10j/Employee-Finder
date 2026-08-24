import os
import re
import json
import asyncio
from pathlib import Path
from docx import Document
from playwright.async_api import async_playwright
from google import genai

OUTPUT_DIR = Path("./candidate_data")
RESUMES_DIR = OUTPUT_DIR / "resumes"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
RESUMES_DIR.mkdir(parents=True, exist_ok=True)

async def extract_candidate_info(file_path: Path, client: genai.Client) -> dict:
    print(f"   🤖 Extracting candidate profile with Gemini: {file_path.name}")
    
    prompt = """
    Analyze this resume document and extract the following candidate details in pure JSON format:
    {
      "full_name": "String or Unknown",
      "dob_or_age": "String or Unknown",
      "total_experience": "String or Unknown",
      "phone_number": "String or Unknown",
      "email": "String or Unknown"
    }
    Return ONLY valid JSON with no markdown block formatting.
    """
    
    try:
        uploaded_file = client.files.upload(file=str(file_path))
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[uploaded_file, prompt]
        )
        raw_text = response.text.strip().replace("```json", "").replace("```", "")
        return json.loads(raw_text)
    except Exception as e:
        print(f"   ⚠️ Gemini extraction note ({file_path.name}): {e}")
        return {
            "full_name": file_path.stem,
            "dob_or_age": "Unknown",
            "total_experience": "Unknown",
            "phone_number": "Unknown",
            "email": "Unknown"
        }

def save_docx_report(candidates: list, output_path: Path):
    doc = Document()
    doc.add_heading("Candidate Sourcing Report", level=0)
    
    for candidate in candidates:
        doc.add_heading(candidate.get("full_name", "Unknown Candidate"), level=1)
        p = doc.add_paragraph()
        p.add_run("Age / DOB: ").bold = True
        p.add_run(f"{candidate.get('dob_or_age')}\n")
        p.add_run("Total Experience: ").bold = True
        p.add_run(f"{candidate.get('total_experience')}\n")
        p.add_run("Phone Number: ").bold = True
        p.add_run(f"{candidate.get('phone_number')}\n")
        p.add_run("Email: ").bold = True
        p.add_run(f"{candidate.get('email')}\n")
        p.add_run("Saved File: ").bold = True
        p.add_run(f"{candidate.get('saved_filename')}\n")
        doc.add_paragraph("-" * 40)
        
    doc.save(str(output_path))
    print(f"✅ Summary DOCX report saved to: {output_path}")

async def run_automation():
    print("=" * 60)
    print("      AUTOMATED JOB SEEKER RESUME FETCHER & PARSER")
    print("=" * 60)
    
    api_key = input("1. Enter your Gemini API Key: ").strip()
    if not api_key:
        print("❌ Error: API Key is required.")
        return

    client = genai.Client(api_key=api_key)
    
    query = input("2. Enter job search query (e.g., Python Developer resumes filetype:pdf Bangalore): ").strip()
    max_results = input("3. Max candidates/links to inspect (Default 5): ").strip()
    max_results = int(max_results) if max_results.isdigit() else 5
    
    print("\n🚀 Starting automated browser execution...\n")
    candidates_list = []
    
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=False)
        context = await browser.new_context(accept_downloads=True)
        page = await context.new_page()
        
        print(f"🔍 Searching Google for: '{query}'")
        await page.goto(f"https://www.google.com/search?q={query}")
        await page.wait_for_timeout(2000)
        
        links = await page.eval_on_selector_all("a h3", "elements => elements.map(e => e.parentElement.href)")
        valid_links = [l for l in links if l and "google.com" not in l][:max_results]
        
        print(f"🔗 Found {len(valid_links)} potential resume links.\n")
        
        for idx, link in enumerate(valid_links, 1):
            print(f"[{idx}/{len(valid_links)}] Processing: {link}")
            try:
                download_path = RESUMES_DIR / f"temp_{idx}.pdf"
                if link.lower().endswith(('.pdf', '.docx', '.png', '.jpg')):
                    download_path = RESUMES_DIR / f"temp_{idx}{Path(link).suffix}"
                    response = await page.request.get(link)
                    with open(download_path, "wb") as f:
                        f.write(await response.body())
                else:
                    async with page.expect_download(timeout=5000) as download_info:
                        await page.goto(link, wait_until="domcontentloaded")
                    download = await download_info.value
                    download_path = RESUMES_DIR / f"temp_{idx}_{download.suggested_filename}"
                    await download.save_as(download_path)
                
                candidate_info = await extract_candidate_info(download_path, client)
                
                raw_name = candidate_info.get("full_name", f"Candidate_{idx}")
                clean_name = re.sub(r'[^a-zA-Z0-9_\- ]', '', raw_name).strip().replace(" ", "_")
                
                final_ext = download_path.suffix if download_path.suffix else ".pdf"
                final_filename = f"{clean_name}{final_ext}"
                final_file_path = RESUMES_DIR / final_filename
                
                if download_path.exists():
                    download_path.rename(final_file_path)
                
                candidate_info["saved_filename"] = final_filename
                candidates_list.append(candidate_info)
                print(f"   🏷️ Saved resume as: {final_filename}\n")
                
            except Exception as e:
                print(f"   ⏩ Link skipped or no direct download found. ({e})\n")
                continue
                
        await browser.close()
        
    if candidates_list:
        json_report = OUTPUT_DIR / "candidates_report.json"
        docx_report = OUTPUT_DIR / "candidates_report.docx"
        
        with open(json_report, "w", encoding="utf-8") as f:
            json.dump(candidates_list, f, indent=2)
            
        print(f"✅ Summary JSON report saved to: {json_report}")
        save_docx_report(candidates_list, docx_report)
    else:
        print("⚠️ No resumes downloaded in this run.")

if __name__ == "__main__":
    asyncio.run(run_automation())
