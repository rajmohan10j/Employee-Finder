import os
import re
import json
import asyncio
import urllib.parse
from pathlib import Path
from docx import Document
from playwright.async_api import async_playwright
from google import genai
from pypdf import PdfReader

OUTPUT_DIR = Path("./candidate_data")
RESUMES_DIR = OUTPUT_DIR / "resumes"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
RESUMES_DIR.mkdir(parents=True, exist_ok=True)

# Guaranteed sample resumes for testing
SAMPLE_RESUME_URLS = [
    "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf"
]

def extract_local_pypdf(file_path: Path) -> dict:
    """Fallback text extractor using local pypdf library if API quota is reached."""
    print(f"   ⚙️ Using local pypdf parser fallback for: {file_path.name}")
    try:
        reader = PdfReader(str(file_path))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
            
        # Basic regex parsing
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
        
        email = email_match.group(0) if email_match else "Not Found"
        phone = phone_match.group(0) if phone_match else "Not Found"
        
        # Derive name from filename or first line
        first_line = text.strip().split('\n')[0] if text.strip() else ""
        name = first_line[:30].strip() if len(first_line) > 2 else file_path.stem
        
        return {
            "full_name": name if name else file_path.stem,
            "dob_or_age": "Unknown",
            "total_experience": "Extracted via Local PDF Reader",
            "phone_number": phone,
            "email": email
        }
    except Exception as e:
        return {
            "full_name": file_path.stem,
            "dob_or_age": "Unknown",
            "total_experience": "Unknown",
            "phone_number": "Unknown",
            "email": "Unknown"
        }

async def extract_candidate_info(file_path: Path, client: genai.Client) -> dict:
    print(f"   🤖 Sending resume to Gemini API...")
    
    prompt = """
    Extract candidate profile details from this resume document in pure JSON format:
    {
      "full_name": "Full Name",
      "dob_or_age": "Age/DOB or Unknown",
      "total_experience": "Years of experience or Unknown",
      "phone_number": "Phone number or Unknown",
      "email": "Email or Unknown"
    }
    Return ONLY valid JSON with no markdown block formatting.
    """
    
    models_to_try = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-2.0-flash']
    
    for model_name in models_to_try:
        try:
            uploaded_file = client.files.upload(file=str(file_path))
            response = client.models.generate_content(
                model=model_name,
                contents=[uploaded_file, prompt]
            )
            raw_text = response.text.strip().replace("```json", "").replace("```", "").strip()
            return json.loads(raw_text)
        except Exception as e:
            if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                print(f"   ⚠️ Gemini API rate limit hit ({model_name}). Trying local parser...")
                break
            elif "404" in str(e):
                continue
            else:
                print(f"   ⚠️ Gemini note: {e}")
                break
                
    # Fallback to local parsing if Gemini API quota is reached
    return extract_local_pypdf(file_path)

def save_docx_report(candidates: list, output_path: Path):
    doc = Document()
    doc.add_heading("Candidate Sourcing Report", level=0)
    
    for c in candidates:
        doc.add_heading(c.get("full_name", "Candidate Profile"), level=1)
        p = doc.add_paragraph()
        p.add_run("Age / DOB: ").bold = True
        p.add_run(f"{c.get('dob_or_age')}\n")
        p.add_run("Total Experience: ").bold = True
        p.add_run(f"{c.get('total_experience')}\n")
        p.add_run("Phone Number: ").bold = True
        p.add_run(f"{c.get('phone_number')}\n")
        p.add_run("Email: ").bold = True
        p.add_run(f"{c.get('email')}\n")
        p.add_run("Saved File: ").bold = True
        p.add_run(f"{c.get('saved_filename')}\n")
        doc.add_paragraph("-" * 40)
        
    doc.save(str(output_path))
    print(f"📄 Word report successfully generated: {output_path}")

async def run_automation():
    print("=" * 65)
    print("      AUTOMATED JOB SEEKER RESUME FETCHER & PARSER v4")
    print("=" * 65)
    
    api_key = input("1. Enter your Gemini API Key: ").strip()
    if not api_key:
        print("❌ Error: API Key is required.")
        return

    client = genai.Client(api_key=api_key)
    
    raw_query = input("2. Enter candidate query (e.g. 'AI Engineer Bangalore'): ").strip()
    if not raw_query:
        raw_query = "AI Engineer Bangalore"
        
    # Auto-enhance query to target PDF resumes
    if "filetype:pdf" not in raw_query.lower():
        search_query = f"{raw_query} resume filetype:pdf"
    else:
        search_query = raw_query
        
    print(f"\n🔍 Optimized Search Query: '{search_query}'")
    print("🚀 Launching browser automation...\n")
    candidates_list = []
    
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=False,
            args=["--disable-blink-features=AutomationControlled"]
        )
        context = await browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120.0.0.0 Safari/537.36",
            accept_downloads=True
        )
        page = await context.new_page()
        
        search_url = f"https://www.google.com/search?q={urllib.parse.quote(search_query)}"
        await page.goto(search_url, wait_until="domcontentloaded")
        await page.wait_for_timeout(3000)
        
        hrefs = await page.eval_on_selector_all("a", "elements => elements.map(e => e.href)")
        
        google_domains = ["google.", "youtube.", "gstatic.", "schema.org"]
        candidate_links = []
        for h in hrefs:
            if h and h.startswith("http"):
                if not any(domain in h.lower() for domain in google_domains):
                    if h not in candidate_links:
                        candidate_links.append(h)
                        
        candidate_links = candidate_links[:5]
        
        if not candidate_links:
            print("⚠️ No direct external resume links found from Google search.")
            print("🔄 Using sample test PDF to demonstrate file download & parsing workflow...")
            candidate_links = SAMPLE_RESUME_URLS
        else:
            print(f"🔗 Found {len(candidate_links)} direct resume links to process.\n")
            
        for idx, link in enumerate(candidate_links, 1):
            print(f"[{idx}/{len(candidate_links)}] Fetching: {link}")
            temp_path = RESUMES_DIR / f"temp_{idx}.pdf"
            
            try:
                response = await page.request.get(link, timeout=12000)
                body = await response.body()
                
                if len(body) < 100:
                    print(f"   ⏩ File empty or skipped.")
                    continue
                    
                with open(temp_path, "wb") as f:
                    f.write(body)
                print(f"   📥 Downloaded file ({len(body)} bytes)")
                
                # Extract candidate details via Gemini API or Local PyPDF
                info = await extract_candidate_info(temp_path, client)
                
                # Sanitize name for file renaming
                raw_name = info.get("full_name", f"Candidate_{idx}")
                clean_name = re.sub(r'[^a-zA-Z0-9_\- ]', '', raw_name).strip().replace(" ", "_")
                if not clean_name or len(clean_name) < 2:
                    clean_name = f"Candidate_{idx}"
                    
                final_filename = f"{clean_name}.pdf"
                final_file_path = RESUMES_DIR / final_filename
                
                # Safe file renaming
                if temp_path.exists():
                    if temp_path.resolve() != final_file_path.resolve():
                        if final_file_path.exists():
                            final_file_path.unlink()
                        temp_path.rename(final_file_path)
                    
                info["saved_filename"] = final_filename
                candidates_list.append(info)
                print(f"   🏷️ Saved resume as: {final_filename}\n")
                
            except Exception as e:
                print(f"   ⏩ Link processing note: {e}\n")
                if temp_path.exists():
                    try:
                        temp_path.unlink()
                    except Exception:
                        pass
                continue
                
        await browser.close()
        
    if candidates_list:
        json_report = OUTPUT_DIR / "candidates_report.json"
        docx_report = OUTPUT_DIR / "candidates_report.docx"
        
        with open(json_report, "w", encoding="utf-8") as f:
            json.dump(candidates_list, f, indent=2)
            
        print(f"✅ JSON report saved to: {json_report}")
        save_docx_report(candidates_list, docx_report)
        print("\n🎉 EXECUTION COMPLETE! Open the 'candidate_data' folder to see your outputs.")
    else:
        print("❌ No candidate files were processed.")

if __name__ == "__main__":
    asyncio.run(run_automation())
