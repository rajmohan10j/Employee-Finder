import os
import re
import json
import asyncio
import urllib.parse
from pathlib import Path
from docx import Document
from playwright.async_api import async_playwright
from google import genai

OUTPUT_DIR = Path("./candidate_data")
RESUMES_DIR = OUTPUT_DIR / "resumes"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
RESUMES_DIR.mkdir(parents=True, exist_ok=True)

# Reliable public resume PDFs for testing or fallback
SAMPLE_RESUME_URLS = [
    "https://raw.githubusercontent.com/AnandChowdhary/resume/master/resume.pdf",
    "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf"
]

async def extract_candidate_info(file_path: Path, client: genai.Client) -> dict:
    print(f"   🤖 Parsing resume with Gemini API: {file_path.name}")
    
    prompt = """
    Extract candidate profile details from this resume document in pure JSON format:
    {
      "full_name": "Full Name or Candidate_X",
      "dob_or_age": "Age/DOB or Unknown",
      "total_experience": "Years of experience or Unknown",
      "phone_number": "Phone number or Unknown",
      "email": "Email or Unknown"
    }
    Return ONLY valid JSON with no markdown block formatting.
    """
    
    # Try models in order of availability
    candidate_models = ['gemini-2.0-flash', 'gemini-1.5-flash']
    
    for model_name in candidate_models:
        try:
            uploaded_file = client.files.upload(file=str(file_path))
            response = client.models.generate_content(
                model=model_name,
                contents=[uploaded_file, prompt]
            )
            raw_text = response.text.strip().replace("```json", "").replace("```", "").strip()
            return json.loads(raw_text)
        except Exception as e:
            if "404" in str(e) or "NOT_FOUND" in str(e):
                continue
            else:
                print(f"   ⚠️ Extraction note ({model_name}): {e}")
                break
                
    return {
        "full_name": file_path.stem,
        "dob_or_age": "Unknown",
        "total_experience": "Unknown",
        "phone_number": "Unknown",
        "email": "Unknown"
    }

def save_docx_report(candidates: list, output_path: Path):
    doc = Document()
    doc.add_heading("Candidate Sourcing Report", level=0)
    
    for c in candidates:
        doc.add_heading(c.get("full_name", "Unknown Candidate"), level=1)
        p = doc.add_paragraph()
        p.add_run("Age / DOB: ").bold = True
        p.add_run(f"{c.get('dob_or_age')}\n")
        p.add_run("Total Experience: ").bold = True
        p.add_run(f"{c.get('total_experience')}\n")
        p.add_run("Phone Number: ").bold = True
        p.add_run(f"{c.get('phone_number')}\n")
        p.add_run("Email: ").bold = True
        p.add_run(f"{c.get('email')}\n")
        p.add_run("Saved File: ").bold = True
        p.add_run(f"{c.get('saved_filename')}\n")
        doc.add_paragraph("-" * 40)
        
    doc.save(str(output_path))
    print(f"📄 DOCX report generated: {output_path}")

async def run_automation():
    print("=" * 65)
    print("      AUTOMATED JOB SEEKER RESUME FETCHER & PARSER v3")
    print("=" * 65)
    
    api_key = input("1. Enter your Gemini API Key: ").strip()
    if not api_key:
        print("❌ Error: API Key is required.")
        return

    client = genai.Client(api_key=api_key)
    
    query = input("2. Enter search query (or press Enter for default 'Software Engineer resume filetype:pdf'): ").strip()
    if not query:
        query = "Software Engineer resume filetype:pdf"
        
    print("\n🚀 Launching browser automation...\n")
    candidates_list = []
    
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=False,
            args=["--disable-blink-features=AutomationControlled"]
        )
        context = await browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120.0.0.0 Safari/537.36",
            accept_downloads=True
        )
        page = await context.new_page()
        
        search_url = f"https://www.google.com/search?q={urllib.parse.quote(query)}"
        print(f"🔍 Searching Google: '{query}'")
        await page.goto(search_url, wait_until="domcontentloaded")
        await page.wait_for_timeout(3000)
        
        hrefs = await page.eval_on_selector_all("a", "elements => elements.map(e => e.href)")
        
        # Strict URL filter to exclude Google system links
        google_domains = ["google.", "youtube.", "gstatic.", "schema.org", "w3.org"]
        candidate_links = []
        for h in hrefs:
            if h and h.startswith("http"):
                if not any(domain in h.lower() for domain in google_domains):
                    if h not in candidate_links:
                        candidate_links.append(h)
                        
        candidate_links = candidate_links[:5]
        
        if not candidate_links:
            print("⚠️ No direct external resume links found from Google search results.")
            print("🔄 Using sample public candidate PDFs to test execution...")
            candidate_links = SAMPLE_RESUME_URLS
        else:
            print(f"🔗 Found {len(candidate_links)} valid candidate links to process.\n")
            
        for idx, link in enumerate(candidate_links, 1):
            print(f"[{idx}/{len(candidate_links)}] Fetching: {link}")
            temp_path = RESUMES_DIR / f"temp_{idx}.pdf"
            
            try:
                response = await page.request.get(link, timeout=12000)
                content_type = response.headers.get("content-type", "")
                
                # Check for PDF/doc/binary payload or force download
                body = await response.body()
                if len(body) < 1000:
                    print(f"   ⏩ Link skipped (response too small/empty).")
                    continue
                    
                with open(temp_path, "wb") as f:
                    f.write(body)
                print(f"   📥 Downloaded file ({len(body)} bytes)")
                
                # Extract candidate details via Gemini API
                info = await extract_candidate_info(temp_path, client)
                
                # Sanitize name for file system
                raw_name = info.get("full_name", f"Candidate_{idx}")
                clean_name = re.sub(r'[^a-zA-Z0-9_\- ]', '', raw_name).strip().replace(" ", "_")
                if not clean_name or len(clean_name) < 2:
                    clean_name = f"Candidate_{idx}"
                    
                final_filename = f"{clean_name}.pdf"
                final_file_path = RESUMES_DIR / final_filename
                
                if temp_path.exists():
                    if final_file_path.exists():
                        final_file_path.unlink()
                    temp_path.rename(final_file_path)
                    
                info["saved_filename"] = final_filename
                candidates_list.append(info)
                print(f"   🏷️ Successfully saved resume as: {final_filename}\n")
                
            except Exception as e:
                print(f"   ⏩ Link processing error: {e}\n")
                if temp_path.exists():
                    temp_path.unlink()
                continue
                
        await browser.close()
        
    if candidates_list:
        json_report = OUTPUT_DIR / "candidates_report.json"
        docx_report = OUTPUT_DIR / "candidates_report.docx"
        
        with open(json_report, "w", encoding="utf-8") as f:
            json.dump(candidates_list, f, indent=2)
            
        print(f"✅ JSON report saved to: {json_report}")
        save_docx_report(candidates_list, docx_report)
        print("\n🎉 EXECUTION SUCCESSFUL! Check the 'candidate_data' folder.")
    else:
        print("❌ No resumes processed in this session.")

if __name__ == "__main__":
    asyncio.run(run_automation())
